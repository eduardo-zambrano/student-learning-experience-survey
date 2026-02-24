#!/usr/bin/env python3
"""
fix_styles.py — Style cleanup for Two_Part_Survey_Instrument_Proposal.docx

Applies consistent formatting (Calibri, 1.15 line spacing, proper heading sizes,
block-quote styling, bullet normalization), removes empty paragraphs, and
applies light language edits for readability.
"""

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import copy
import os

FILEPATH = os.path.join(os.path.dirname(__file__),
                        "Two_Part_Survey_Instrument_Proposal.docx")


def configure_style(style, font_name="Calibri", font_size=Pt(11),
                     bold=False, italic=False, color=None,
                     space_before=None, space_after=Pt(6),
                     line_spacing=1.15, left_indent=None,
                     alignment=None):
    """Set all attributes on a style object."""
    font = style.font
    font.name = font_name
    font.size = font_size
    font.bold = bold
    font.italic = italic
    font.color.rgb = color if color else RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.left_indent = left_indent
    pf.alignment = alignment


def set_global_styles(doc):
    """Define the five global styles."""
    styles = doc.styles

    # --- Normal ---
    configure_style(styles['Normal'],
                    font_size=Pt(11), space_after=Pt(6))

    # --- Heading 1 ---
    configure_style(styles['Heading 1'],
                    font_size=Pt(16), bold=True,
                    space_before=Pt(12), space_after=Pt(6))

    # --- Heading 2 ---
    configure_style(styles['Heading 2'],
                    font_size=Pt(13), bold=True,
                    space_before=Pt(12), space_after=Pt(6))

    # --- Quote ---
    configure_style(styles['Quote'],
                    font_size=Pt(10), italic=True,
                    left_indent=Inches(0.5), space_after=Pt(6))
    # Remove italic from char style too if present, we'll control via runs
    if styles['Quote'].font:
        styles['Quote'].font.italic = True

    # --- List Bullet ---
    configure_style(styles['List Bullet'],
                    font_size=Pt(11),
                    left_indent=Inches(0.35), space_after=Pt(3))


def remove_paragraph(paragraph):
    """Delete a paragraph element from the document XML tree."""
    elem = paragraph._element
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)


def fix_typo(paragraph):
    """Fix 'Berlekey' → 'Berkeley' in paragraph runs."""
    for run in paragraph.runs:
        if "Berlekey" in run.text:
            run.text = run.text.replace("Berlekey", "Berkeley")


def clear_run_formatting(paragraph):
    """Clear run-level font overrides so the style controls appearance."""
    for run in paragraph.runs:
        run.font.name = None
        run.font.size = None
        run.font.bold = None
        run.font.italic = None
        run.font.color.rgb = None


def clear_paragraph_direct_formatting(paragraph):
    """Clear paragraph-level direct formatting so style takes over."""
    pf = paragraph.paragraph_format
    pf.space_before = None
    pf.space_after = None
    pf.line_spacing = None
    # Don't clear left_indent or alignment by default — handled per case


def replace_in_runs(paragraph, old, new):
    """Find-and-replace across run boundaries in a paragraph.

    Concatenates all run texts, performs the replacement, then redistributes
    the result back across the original runs (preserving their formatting).
    """
    runs = paragraph.runs
    if not runs:
        return
    combined = "".join(r.text for r in runs)
    if old not in combined:
        return
    updated = combined.replace(old, new)
    # Redistribute: fill runs in order, empty leftovers
    pos = 0
    for r in runs:
        end = pos + len(r.text)
        if pos < len(updated):
            r.text = updated[pos:min(end + (len(updated) - len(combined)), len(updated))]
        else:
            r.text = ""
        pos = end
    # Simpler approach: put all text in first run, clear the rest
    # (This is safer when lengths change significantly)
    for r in runs:
        r.text = ""
    runs[0].text = updated


def replace_in_runs_preserving_structure(paragraph, old, new):
    """Find-and-replace within individual runs only (no cross-run merging).

    Use this when the target text is known to be within a single run.
    """
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def find_paragraph(paragraphs, snippet):
    """Return the first paragraph whose text contains the given snippet."""
    for p in paragraphs:
        if snippet in p.text:
            return p
    return None


def apply_language_edits(doc):
    """Apply nine light language edits for readability.

    Edits match paragraphs by text content (not index) so the script is
    idempotent and safe to re-run.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraphs = doc.paragraphs

    # --- Edit 1: Remove redundant "proposal" ---
    p = find_paragraph(paragraphs,
                       "student survey instrument proposal composed")
    if p:
        for run in p.runs:
            if run.text.strip() == "proposal":
                run.text = ""
                break
        print("  Edit 1: removed redundant 'proposal'")
    else:
        print("  Edit 1: already applied (no match)")

    # --- Edit 2: Summative — rephrase + Oxford comma ---
    p = find_paragraph(paragraphs,
                       "must not have any questionnaire items that have been "
                       "clearly documented in the literature to contain biases")
    if p:
        replace_in_runs_preserving_structure(
            p,
            "must not have any questionnaire items that have been clearly "
            "documented in the literature to contain biases against "
            "individuals in protected categories, including but not limited "
            "to gender, ethnicity and age",
            "must not include any questionnaire items that the literature "
            "has clearly documented as biased against individuals in "
            "protected categories, including but not limited to gender, "
            "ethnicity, and age"
        )
        print("  Edit 2: summative rephrase + Oxford comma")
    else:
        print("  Edit 2: already applied (no match)")

    # --- Edit 3: Formative — rephrase "items…items" + Oxford comma ---
    p = find_paragraph(paragraphs,
                       "to contain items that the students are not qualified")
    if p:
        replace_in_runs(
            p,
            "component must not have any questionnaire items that have "
            "been clearly documented in the literature to contain items "
            "that the students are not qualified to assess, including but "
            "not limited to disciplinary expertise, pedagogical "
            "effectiveness and learning outcome attainment",
            "component must not include questionnaire items on topics that "
            "the literature has clearly documented as beyond what students "
            "are qualified to assess, including but not limited to "
            "disciplinary expertise, pedagogical effectiveness, and "
            "learning outcome attainment"
        )
        print("  Edit 3: formative rephrase + Oxford comma")
    else:
        print("  Edit 3: already applied (no match)")

    # --- Edit 4: Empty bullet paragraph — handled by removal logic ---
    print("  Edit 4: empty bullet (handled by removal)")

    # --- Edit 5: "working progress" → "work in progress",
    #     "tending to" → "following", remove comma before "and" ---
    p = find_paragraph(paragraphs, "working progress")
    if p:
        for run in p.runs:
            if "working progress" in run.text:
                run.text = run.text.replace(
                    "working progress", "work in progress")
        print("  Edit 5a: 'work in progress'")
    else:
        print("  Edit 5a: already applied (no match)")

    p = find_paragraph(paragraphs, "tending to the criteria")
    if p:
        for run in p.runs:
            if "tending to" in run.text:
                run.text = run.text.replace(
                    "were drafted tending to the criteria in",
                    "were drafted following the criteria in")
        print("  Edit 5b: 'following'")
    else:
        print("  Edit 5b: already applied (no match)")

    # Remove comma before "and the research" (XML-level for hyperlink para)
    p = find_paragraph(paragraphs, ", and the research by Berkeley")
    if p:
        p_elem = p._element
        for r_elem in p_elem.findall(f"{{{W}}}r"):
            t_elem = r_elem.find(f"{{{W}}}t")
            if (t_elem is not None and t_elem.text
                    and ", and the research" in t_elem.text):
                t_elem.text = t_elem.text.replace(
                    ", and the research", " and the research")
        print("  Edit 5c: removed comma before 'and'")
    else:
        print("  Edit 5c: already applied (no match)")

    # --- Edit 6: Empty paragraph — handled by removal logic ---
    print("  Edit 6: empty paragraph (handled by removal)")

    # --- Edit 7: Add comma before "as" in scoring paragraph ---
    p = find_paragraph(paragraphs,
                       "numerical values as those values cannot")
    if p:
        replace_in_runs_preserving_structure(
            p, " as those values", ", as those values")
        print("  Edit 7: added comma before 'as'")
    else:
        print("  Edit 7: already applied (no match)")

    # --- Edit 8: Remove "as follows:" + add "the" before "total" ---
    p = find_paragraph(paragraphs,
                       "counts as follows: for example")
    if p:
        replace_in_runs_preserving_structure(
            p,
            "ratios of counts as follows: for example, how many out of "
            "total",
            "ratios of counts: for example, how many out of the total"
        )
        print("  Edit 8: removed 'as follows', added 'the'")
    else:
        print("  Edit 8: already applied (no match)")

    # --- Edit 9: Empty paragraph — handled by removal logic ---
    print("  Edit 9: empty paragraph (handled by removal)")


def replace_question(paragraph, new_label, new_text):
    """Replace a survey-question paragraph's label and question text.

    Expects the standard 3-run structure:
      R0 = "N. "  (number — kept unchanged)
      R1 = bold label
      R2 = ': "question text"'
    Preserves run formatting (bold on R1, normal on R0/R2).
    """
    runs = paragraph.runs
    if len(runs) < 3:
        # Fallback: put everything in a single run
        combined = paragraph.text
        number = combined.split(".")[0] + ". "
        for r in runs:
            r.text = ""
        runs[0].text = f'{number}{new_label}: {new_text}'
        return
    runs[1].text = new_label
    runs[2].text = f': {new_text}'


def apply_question_replacements(doc):
    """Replace the six survey questions with the revised instrument.

    Summative:
      S1 (participatory climate)  — unchanged
      S2 (approachability)        — unchanged
      S3 "Communication & clarity" → "Respectful treatment"

    Formative:
      F1 "Learning enhancement"  → "Engagement"
      F2 "Assignment design"     → "Workload"
      F3 "Inclusive environment"  → "Interest change"

    Also updates the commentary paragraph (P19) and source line (P25).
    Matches paragraphs by text content for idempotency.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraphs = doc.paragraphs

    # ── S3: "Communication & clarity" → "Respectful treatment" ──
    p = find_paragraph(paragraphs, "Communication & clarity")
    if p:
        replace_question(
            p,
            "Respectful treatment",
            '\u201cI felt that the instructor treated students '
            'with respect.\u201d'
        )
        print("  Q-S3: replaced 'Communication & clarity' → "
              "'Respectful treatment'")
    else:
        # Check if already applied
        if find_paragraph(paragraphs, "Respectful treatment"):
            print("  Q-S3: already applied")
        else:
            print("  Q-S3: WARNING — paragraph not found")

    # ── P19: Update commentary paragraph ──
    # Old: "These questions (a work in progress for the time being)
    #        were drafted following the criteria in the presentation
    #        by Committee Member Jeff Palmer and the research by
    #        Berkeley Statistics Professor Philip Stark."
    # New: Remove "work in progress" hedge, keep attribution to
    #      Palmer presentation and Stark research.
    p = find_paragraph(paragraphs,
                       "work in progress for the time being")
    if p:
        # This paragraph has a complex run/hyperlink structure.
        # Replace "a work in progress for the time being" with
        # the updated text using XML-level manipulation.
        p_elem = p._element
        for r_elem in p_elem.findall(f"{{{W}}}r"):
            t_elem = r_elem.find(f"{{{W}}}t")
            if t_elem is not None and t_elem.text:
                if "These questions " in t_elem.text:
                    t_elem.text = "These questions "
                elif t_elem.text == "(":
                    t_elem.text = ""
                elif "a work in progress for the time being" in t_elem.text:
                    t_elem.text = ""
                elif t_elem.text.startswith(") "):
                    t_elem.text = ""
        print("  Q-P19: removed 'work in progress' hedge")
    else:
        if find_paragraph(paragraphs, "These questions were drafted"):
            print("  Q-P19: already applied")
        else:
            print("  Q-P19: WARNING — paragraph not found")

    # ── F1: "Learning enhancement" → "Engagement" ──
    p = find_paragraph(paragraphs, "Learning enhancement")
    if p:
        replace_question(
            p,
            "Engagement",
            '\u201cI found the class sessions engaging.\u201d'
        )
        print("  Q-F1: replaced 'Learning enhancement' → 'Engagement'")
    else:
        if find_paragraph(paragraphs, "Engagement"):
            print("  Q-F1: already applied")
        else:
            print("  Q-F1: WARNING — paragraph not found")

    # ── F2: "Assignment design" → "Workload" ──
    p = find_paragraph(paragraphs, "Assignment design")
    if p:
        replace_question(
            p,
            "Workload",
            '\u201cThe workload in this course was manageable '
            'given the number of units.\u201d'
        )
        print("  Q-F2: replaced 'Assignment design' → 'Workload'")
    else:
        if find_paragraph(paragraphs, "Workload"):
            print("  Q-F2: already applied")
        else:
            print("  Q-F2: WARNING — paragraph not found")

    # ── F3: "Inclusive environment" → "Interest change" ──
    p = find_paragraph(paragraphs, "Inclusive environment")
    if p:
        replace_question(
            p,
            "Interest change",
            '\u201cThis course increased my interest in the '
            'subject.\u201d'
        )
        print("  Q-F3: replaced 'Inclusive environment' → "
              "'Interest change'")
    else:
        if find_paragraph(paragraphs, "Interest change"):
            print("  Q-F3: already applied")
        else:
            print("  Q-F3: WARNING — paragraph not found")

    # ── P25: Update source line ──
    # Old: "Source: UC Berkeley Academic Senate, DIVCO to VPF, …"
    # New: committee-designed questions; update attribution.
    p = find_paragraph(paragraphs,
                       "Source: UC Berkeley Academic Senate")
    if p:
        # Remove hyperlink and replace with updated source text
        p_elem = p._element
        # Remove all hyperlinks from this paragraph
        for h in p_elem.findall(f"{{{W}}}hyperlink"):
            p_elem.remove(h)
        # Clear existing runs and set new text
        for r in p.runs:
            r.text = ""
        p.runs[0].text = (
            "Source: Committee-designed questions informed by "
            "Philip Stark\u2019s expert report (Appendix\u00a0A) "
            "and the committee\u2019s adopted criteria."
        )
        print("  Q-P25: updated source line")
    else:
        if find_paragraph(paragraphs, "Committee-designed questions"):
            print("  Q-P25: already applied")
        else:
            print("  Q-P25: WARNING — paragraph not found")


def main():
    doc = Document(FILEPATH)

    # ── Step 0a: Apply language edits (before any style/removal work) ──
    print("Applying language edits...")
    apply_language_edits(doc)

    # ── Step 0b: Apply question replacements ──
    print("Applying question replacements...")
    apply_question_replacements(doc)

    # ── Step 1: Set global style definitions ──
    set_global_styles(doc)

    # ── Step 2: Collect paragraphs and identify which to remove ──
    paragraphs = list(doc.paragraphs)
    total = len(paragraphs)

    # Build list of paragraph indices to remove (empty / meaningless ones)
    to_remove = set()
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            # Remove all empty paragraphs
            to_remove.add(i)

    # ── Step 3: Process each paragraph (before removal, since indices matter) ──
    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # --- Fix typo in P20 ---
        if "Berlekey" in p.text:
            fix_typo(p)

        # --- Skip paragraphs marked for removal ---
        if i in to_remove:
            continue

        # --- Appendix A heading (P52): List Bullet → Heading 2 ---
        if text.startswith("Appendix A:"):
            p.style = doc.styles['Heading 2']
            p.paragraph_format.left_indent = None  # no indent for headings
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- Appendix B heading (P61): List Bullet → Heading 2 ---
        if text.startswith("Appendix B:"):
            p.style = doc.styles['Heading 2']
            p.paragraph_format.left_indent = None
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- Stark passage block quote (P56) ---
        if text.startswith("(…) Items relating to teaching effectiveness"):
            p.style = doc.styles['Quote']
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.name = None  # inherit Calibri from style
                run.font.bold = None
            clear_paragraph_direct_formatting(p)
            continue

        # --- Source line after Stark quote (P58) → Normal ---
        if text.startswith("Source: https://www.tfanet.ca"):
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = None
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- Appendix B checklist items: normalize indent ---
        if p.style.name == 'List Bullet' and text and i > 50:
            # These are the checklist items in Appendix B area
            p.paragraph_format.left_indent = Inches(0.35)
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- All other List Bullet paragraphs: normalize indent ---
        if p.style.name == 'List Bullet':
            indent = p.paragraph_format.left_indent
            # Fix 228600 EMU (0.18in) or 0 or None → 0.35in
            if indent is None or indent == 0 or indent == Emu(228600):
                p.paragraph_format.left_indent = Inches(0.35)
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- Heading paragraphs: clear run-level overrides ---
        if p.style.name in ('Heading 1', 'Heading 2'):
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            # Keep center alignment on title (P00)
            if i == 0:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.paragraph_format.alignment = None
            continue

        # --- Quote paragraphs (Background section) ---
        if p.style.name == 'Quote':
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            continue

        # --- Normal paragraphs: clear run-level overrides ---
        if p.style.name == 'Normal':
            clear_run_formatting(p)
            clear_paragraph_direct_formatting(p)
            # Subtitle line (P01) stays centered
            if i == 1:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            # Survey question items (indented Normal paragraphs)
            elif p.paragraph_format.left_indent and p.paragraph_format.left_indent > Emu(300000):
                # Re-apply bold on the label run (R1)
                if len(p.runs) >= 2:
                    p.runs[1].font.bold = True
            continue

    # ── Step 4: Remove empty paragraphs (iterate in reverse to preserve indices) ──
    for i in sorted(to_remove, reverse=True):
        remove_paragraph(paragraphs[i])

    # ── Step 5: Save ──
    doc.save(FILEPATH)
    print(f"Saved: {FILEPATH}")

    # ── Step 6: Verification dump ──
    doc2 = Document(FILEPATH)
    print(f"\n{'='*80}")
    print(f"VERIFICATION: {len(doc2.paragraphs)} paragraphs after cleanup")
    print(f"{'='*80}")
    for i, p in enumerate(doc2.paragraphs):
        pf = p.paragraph_format
        txt = p.text[:90].replace('\n', ' ') if p.text else ''
        indent = pf.left_indent
        print(f"P{i:02d} [{p.style.name}] indent={indent} "
              f"sb={pf.space_before} sa={pf.space_after} "
              f"\"{txt}\"")


if __name__ == "__main__":
    main()
