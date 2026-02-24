"""
Generate the Two-Part Survey Instrument Proposal as a .docx file.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global style defaults ──────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Adjust heading styles
for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.name = "Calibri"

doc.styles["Heading 1"].font.size = Pt(16)
doc.styles["Heading 2"].font.size = Pt(13)
doc.styles["Heading 3"].font.size = Pt(11)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ── Helper functions ───────────────────────────────────────────────
def add_question_block(doc, number, label, text):
    """Add a formatted question with label and italic block-quote text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    run_num = p.add_run(f"{number}. ")
    run_num.bold = True

    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.italic = True

    run_text = p.add_run(f"\u201c{text}\u201d")
    run_text.italic = True


def add_bullet(doc, text, bold_prefix=None, indent=0.35):
    """Add a bullet-point paragraph, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    return p


# ── Title ──────────────────────────────────────────────────────────
title = doc.add_heading("Proposal for a Two-Part Student Survey Instrument", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Ad Hoc Committee on Student Perceptions of Teaching Effectiveness\n"
    "Orfalea College of Business, Cal Poly"
)
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(80, 80, 80)

# ── 1. Introduction / Rationale ───────────────────────────────────
doc.add_heading("1. Introduction / Rationale", level=2)

add_body(
    doc,
    "This proposal presents a student survey instrument composed of two distinct parts\u2014"
    "a summative component and a formative component\u2014each serving a different purpose "
    "in the evaluation of teaching."
)

add_bullet(
    doc,
    "The summative component addresses what the University Faculty Personnel Actions (UFPA) "
    "document requires regarding student evaluations (Section 7.2.5.4). Specifically, it helps "
    "evaluators assess Factor 8 (\u201crelationship with students in class\u201d) and related observable "
    "student experiences.",
    bold_prefix="Summative component. ",
)

add_bullet(
    doc,
    "The summative component is the only part of the survey that is placed in the "
    "Personnel Action File (PAF).",
    bold_prefix="PAF inclusion. ",
)

add_bullet(
    doc,
    "The formative component provides developmental feedback directly to the instructor "
    "and does not go to the PAF.",
    bold_prefix="Formative component. ",
)

add_bullet(
    doc,
    "The committee recommends the use of a rubric for the evaluation of the other nine "
    "teaching performance factors per the AP-109 form. Development of that rubric is "
    "outside the scope of the current proposal.",
    bold_prefix="Rubric for other factors. ",
)

# ── 2. Summative Part ─────────────────────────────────────────────
doc.add_heading("2. Summative Part (Three Questions \u2014 Goes to PAF)", level=2)

add_body(
    doc,
    "The following three questions constitute the summative portion of the survey. "
    "This is the only part that is placed in the Personnel Action File (PAF). "
    "It is designed exclusively to help the evaluator assess Factor 8 and related "
    "student-observable aspects of teaching per the UFPA."
)

add_question_block(
    doc,
    1,
    "Classroom experience & climate",
    "I felt comfortable participating in class (asking questions, sharing ideas, "
    "or making mistakes).",
)

add_question_block(
    doc,
    2,
    "Classroom experience & climate / approachability",
    "I felt the instructor was approachable if I needed help (in office hours, "
    "after class, or by email).",
)

add_question_block(
    doc,
    3,
    "Communication & clarity",
    "Expectations for assignments and grading were communicated clearly "
    "(e.g., instructions, rubrics, deadlines, examples).",
)

source_p = doc.add_paragraph()
source_p.paragraph_format.space_before = Pt(8)
run = source_p.add_run("Source: ")
run.bold = True
source_p.add_run("Committee presentation materials and deliberations.")

# ── 3. Formative Part ─────────────────────────────────────────────
doc.add_heading("3. Formative Part (Three Questions \u2014 Does NOT Go to PAF)", level=2)

add_body(
    doc,
    "The following three questions constitute the formative portion of the survey. "
    "Responses are shared only with the instructor for developmental purposes and "
    "do not go to the PAF."
)

add_question_block(
    doc,
    1,
    "Learning enhancement",
    "The instructor\u2019s lectures, facilitation of classes, and/or office hours and "
    "help sessions enhanced my learning. (\u2018Learning\u2019 may include gaining mastery of "
    "course content and new skills, exposure to new methodologies and modes of "
    "critical thinking, and extending the ability to express oneself on the topics "
    "treated in the course.)",
)

add_question_block(
    doc,
    2,
    "Assignment design",
    "The assignments were well designed to help me understand the course material "
    "and gain a deeper perspective on the subject.",
)

add_question_block(
    doc,
    3,
    "Inclusive environment",
    "The instructor created an environment in which I could feel included "
    "(for example, encouraged multiple voices/perspectives, welcomed questions "
    "and critiques, responded to student feedback).",
)

source_p = doc.add_paragraph()
source_p.paragraph_format.space_before = Pt(8)
run = source_p.add_run("Source: ")
run.bold = True
source_p.add_run(
    "UC Berkeley Academic Senate, DIVCO to VPF, "
    "\u201cEvaluation of Teaching,\u201d June 29, 2021."
)

# ── 4. Scoring Methodology ────────────────────────────────────────
doc.add_heading("4. Scoring Methodology (Applies to Both Parts)", level=2)

add_body(
    doc,
    "The following scoring approach applies to both the summative and formative "
    "portions of the survey."
)

add_bullet(
    doc,
    "Strongly Agree, Agree, Neither Agree nor Disagree, Disagree, Strongly Disagree.",
    bold_prefix="Five ordered categorical response options: ",
)

add_bullet(
    doc,
    "is also available for each question.",
    bold_prefix="A Not Applicable (N/A) option ",
)

add_bullet(
    doc,
    "The categorical responses are not assigned numerical values. "
    "They remain ordered categories.",
    bold_prefix="No numerical scoring. ",
)

add_bullet(
    doc,
    'The resulting data are not to be called "quantitative." '
    "They are ordered categorical data.",
    bold_prefix="Not quantitative. ",
)

# ── 5. Reporting Guidelines ───────────────────────────────────────
doc.add_heading("5. Reporting Guidelines (Applies to Both Parts)", level=2)

add_body(
    doc,
    "The following reporting principles apply to both the summative and formative "
    "portions of the survey."
)

add_bullet(
    doc,
    "Results are presented as ratios of counts: for example, how many out of total "
    "respondents agree or strongly agree, and what fraction of those strongly agree.",
    bold_prefix="Ratios of counts. ",
)

add_bullet(
    doc,
    "The percentage of students whose response falls in each category should be reported.",
    bold_prefix="Frequency distributions. ",
)

add_bullet(
    doc,
    "The proportion of enrolled students who responded should be reported.",
    bold_prefix="Response rates. ",
)

add_bullet(
    doc,
    "Results should not be extrapolated from responders to nonresponders.",
    bold_prefix="No extrapolation. ",
)

add_bullet(
    doc,
    "Results should not be compared across course formats, levels, topics, or "
    "disciplines\u2014for either the formative or the summative responses.",
    bold_prefix="No cross-comparisons. ",
)

# ── 6. Rubric for Other Factors ───────────────────────────────────
doc.add_heading("6. Rubric for Other Factors", level=2)

add_body(
    doc,
    "The committee recommends the use of a rubric for the evaluation of the other "
    "nine teaching performance factors per the AP-109 form. Development of that rubric "
    "is outside the scope of the current proposal but is identified as a priority "
    "for future work."
)

# ── References ─────────────────────────────────────────────────────
doc.add_heading("References", level=2)

refs = [
    "California Polytechnic State University. University Faculty Personnel Actions (UFPA), "
    "Section 7.2.5 (7.2.5.1 through 7.2.5.4).",
    "UC Berkeley Academic Senate, DIVCO to VPF. \u201cEvaluation of Teaching.\u201d June 29, 2021. "
    "https://academic-senate.berkeley.edu/sites/default/files/"
    "divco_to_vpf-eval_of_teaching_w.enc-06.29.2021.pdf",
    "Committee presentation materials and deliberations.",
]

for ref in refs:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(ref)

# ── Save ───────────────────────────────────────────────────────────
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "Two_Part_Survey_Instrument_Proposal.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
